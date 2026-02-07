import argparse
import importlib.util
import json
import logging
import os
import re
import hashlib
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv"}

DEFAULT_CHUNK_TOKENS = 800
DEFAULT_OVERLAP_TOKENS = 120

# Excel-specific chunking: row-aware chunks
DEFAULT_EXCEL_ROWS_PER_CHUNK = 120
DEFAULT_EXCEL_OVERLAP_ROWS = 15


def module_available(module_name: str) -> bool:
    return importlib.util.find_spec(module_name) is not None


def sha1_text(value: str) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()


def normalize_modified_time(path: Path) -> str:
    modified_ts = path.stat().st_mtime
    return datetime.fromtimestamp(modified_ts, tz=timezone.utc).isoformat()


def normalize_for_hash(text: str) -> str:
    # Stable hashing across platforms/whitespace differences.
    return " ".join(text.split())


def chunk_text(text: str, chunk_tokens: int, overlap_tokens: int) -> list[str]:
    """Whitespace-token chunking for unstructured text."""
    tokens = re.findall(r"\S+", text)
    if not tokens:
        return []
    if chunk_tokens <= 0:
        return [" ".join(tokens)]
    if overlap_tokens >= chunk_tokens:
        overlap_tokens = max(chunk_tokens - 1, 0)
    chunks = []
    start = 0
    step = max(chunk_tokens - overlap_tokens, 1)
    while start < len(tokens):
        end = min(start + chunk_tokens, len(tokens))
        chunks.append(" ".join(tokens[start:end]))
        if end == len(tokens):
            break
        start += step
    return chunks


def dataframe_to_text(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a text representation suitable for embedding.
    Prefer markdown table if tabulate is available; otherwise fall back to CSV.
    """
    # Make sure we don't embed NaNs as "nan"
    df = df.fillna("")
    if module_available("tabulate"):
        try:
            return df.to_markdown(index=False)
        except Exception:
            pass
    return df.to_csv(index=False)


def extract_pdf(path: Path) -> list[dict]:
    if module_available("pdfplumber"):
        import pdfplumber

        pages = []
        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append({"text": text, "page_number": index})
        return pages
    if module_available("pypdf"):
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append({"text": text, "page_number": index})
        return pages
    raise RuntimeError("No PDF parser available. Install pdfplumber or pypdf.")


def extract_docx(path: Path) -> list[dict]:
    if not module_available("docx"):
        raise RuntimeError("python-docx is required for .docx files.")
    from docx import Document

    document = Document(str(path))
    lines = [para.text for para in document.paragraphs if para.text]
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                lines.append(" | ".join(row_text))
    return [{"text": "\n".join(lines)}]


def extract_excel_sheets(path: Path) -> list[dict]:
    """
    Returns a list of sections, one per sheet, including the sheet DataFrame.
    Actual chunking happens later in build_manifest_entries as row-aware chunks.
    """
    entries: list[dict] = []
    excel = pd.ExcelFile(path)
    for sheet_name in excel.sheet_names:
        df = excel.parse(sheet_name=sheet_name)

        # Drop completely empty rows/cols to reduce noise
        df = df.dropna(how="all")
        df = df.dropna(axis=1, how="all")

        # If the sheet is entirely empty after cleanup, skip it
        if df.empty:
            continue

        entries.append({"df": df, "sheet_name": sheet_name})
    return entries


def extract_csv(path: Path) -> list[dict]:
    df = pd.read_csv(path)
    df = df.dropna(how="all").dropna(axis=1, how="all")
    return [{"df": df}]


def extract_text(path: Path) -> list[dict]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return extract_excel_sheets(path)
    if suffix == ".csv":
        return extract_csv(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def scan_documents(docs_dir: Path) -> list[Path]:
    return [
        path
        for path in docs_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]


def build_doc_id(source_rel: str, section_tag: str) -> str:
    """
    Stable doc_id across machines:
    - source_rel is the relative path to docs_dir (posix style)
    - section_tag identifies a sheet/page/body
    """
    return sha1_text(f"{source_rel}::{section_tag}")


def build_manifest_entries(
    path: Path,
    docs_dir: Path,
    chunk_tokens: int,
    overlap_tokens: int,
    excel_rows_per_chunk: int,
    excel_overlap_rows: int,
) -> list[dict]:
    entries: list[dict] = []
    try:
        extracted_sections = extract_text(path)
    except Exception as exc:
        logging.warning("Skipping %s due to extraction error: %s", path, exc)
        return entries

    # Stable relative path to avoid Windows/Ubuntu absolute-path churn
    try:
        source_rel = path.resolve().relative_to(docs_dir.resolve()).as_posix()
    except Exception:
        # Fallback if the file isn't under docs_dir (shouldn't happen normally)
        source_rel = path.name

    file_metadata = {
        "source_name": path.name,
        "source_rel": source_rel,          # stable across environments
        "source_path": str(path.resolve()), # keep absolute path for debugging
        "modified_time": normalize_modified_time(path),
        "file_type": path.suffix.lower().lstrip("."),
    }

    for section in extracted_sections:
        # Section tags for stable IDs + filtering
        if "sheet_name" in section:
            section_tag = f"sheet::{section['sheet_name']}"
        elif "page_number" in section:
            section_tag = f"page::{section['page_number']}"
        else:
            section_tag = "body"

        # Excel/CSV row-aware chunking
        if "df" in section:
            df: pd.DataFrame = section["df"]
            sheet_name = section.get("sheet_name")
            n_rows = len(df)

            if excel_rows_per_chunk <= 0:
                excel_rows_per_chunk = n_rows or 1
            if excel_overlap_rows >= excel_rows_per_chunk:
                excel_overlap_rows = max(excel_rows_per_chunk - 1, 0)

            step = max(excel_rows_per_chunk - excel_overlap_rows, 1)

            doc_id = build_doc_id(source_rel, section_tag)

            chunk_index = 0
            row_start = 0
            while row_start < n_rows:
                row_end = min(row_start + excel_rows_per_chunk, n_rows)
                sub = df.iloc[row_start:row_end].copy()
                chunk_text_repr = dataframe_to_text(sub)

                if chunk_text_repr.strip():
                    # Use 1-based row ranges for readability
                    metadata = file_metadata | {
                        "doc_id": doc_id,
                        "chunk_index": chunk_index,
                        "section_tag": section_tag,
                        "sheet_name": sheet_name,
                        "row_start": row_start + 1,
                        "row_end": row_end,
                        "n_rows": n_rows,
                        "n_cols": int(sub.shape[1]),
                    }

                    chunk_id = f"{doc_id}::{chunk_index}"
                    content_hash = sha1_text(normalize_for_hash(chunk_text_repr))

                    entries.append(
                        {
                            "doc_id": doc_id,
                            "chunk_id": chunk_id,
                            "text": chunk_text_repr,
                            "content_hash": content_hash,
                            "metadata": metadata,
                        }
                    )

                    chunk_index += 1

                if row_end == n_rows:
                    break
                row_start += step

            # Add total_chunks after building all for this section (optional)
            # (We keep it in metadata at insertion time for simplicity)
            continue

        # Unstructured text chunking (PDF pages, docx body, etc.)
        raw_text = section.get("text", "")
        if not raw_text.strip():
            logging.warning("No text extracted for %s (metadata: %s)", path, section)
            continue

        doc_id = build_doc_id(source_rel, section_tag)
        chunks = chunk_text(raw_text, chunk_tokens, overlap_tokens)
        for idx, chunk in enumerate(chunks):
            metadata: dict[str, Any] = file_metadata | {
                "doc_id": doc_id,
                "chunk_index": idx,
                "total_chunks": len(chunks),
                "section_tag": section_tag,
            }
            for key in ("sheet_name", "page_number"):
                if key in section:
                    metadata[key] = section[key]

            chunk_id = f"{doc_id}::{idx}"
            content_hash = sha1_text(normalize_for_hash(chunk))

            entries.append(
                {
                    "doc_id": doc_id,
                    "chunk_id": chunk_id,
                    "text": chunk,
                    "content_hash": content_hash,
                    "metadata": metadata,
                }
            )

    return entries


def write_manifest(entries: list[dict], manifest_path: Path) -> None:
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    with manifest_path.open("w", encoding="utf-8") as handle:
        for entry in entries:
            handle.write(json.dumps(entry, ensure_ascii=False) + "\n")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="RAG ingestion manifest builder (pgvector-friendly).")
    parser.add_argument(
        "--docs-dir",
        default=os.environ.get("RAG_DOCS_DIR", "data"),
        help="Directory to scan for documents (env: RAG_DOCS_DIR).",
    )
    parser.add_argument(
        "--manifest-path",
        default=os.environ.get("RAG_MANIFEST_PATH", "rag_manifest.jsonl"),
        help="Output JSONL manifest path (env: RAG_MANIFEST_PATH).",
    )
    parser.add_argument(
        "--chunk-tokens",
        type=int,
        default=DEFAULT_CHUNK_TOKENS,
        help="Approximate tokens per chunk for unstructured text (PDF/DOCX).",
    )
    parser.add_argument(
        "--overlap-tokens",
        type=int,
        default=DEFAULT_OVERLAP_TOKENS,
        help="Approximate overlap tokens between unstructured chunks.",
    )
    parser.add_argument(
        "--excel-rows-per-chunk",
        type=int,
        default=DEFAULT_EXCEL_ROWS_PER_CHUNK,
        help="Rows per chunk for Excel/CSV (row-aware chunking).",
    )
    parser.add_argument(
        "--excel-overlap-rows",
        type=int,
        default=DEFAULT_EXCEL_OVERLAP_ROWS,
        help="Overlapping rows between Excel/CSV chunks.",
    )
    return parser.parse_args()


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    args = parse_args()
    docs_dir = Path(args.docs_dir)
    if not docs_dir.exists():
        raise SystemExit(f"Docs directory not found: {docs_dir}")

    logging.info("Scanning %s for RAG documents.", docs_dir)
    files = scan_documents(docs_dir)
    logging.info("Found %d supported documents.", len(files))

    all_entries: list[dict] = []
    for path in files:
        all_entries.extend(
            build_manifest_entries(
                path=path,
                docs_dir=docs_dir,
                chunk_tokens=args.chunk_tokens,
                overlap_tokens=args.overlap_tokens,
                excel_rows_per_chunk=args.excel_rows_per_chunk,
                excel_overlap_rows=args.excel_overlap_rows,
            )
        )

    write_manifest(all_entries, Path(args.manifest_path))
    logging.info("Wrote %d chunks to %s.", len(all_entries), args.manifest_path)


if __name__ == "__main__":
    main()
